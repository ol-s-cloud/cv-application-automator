import re
import os
from docx import Document
from docx2pdf import convert
from pathlib import Path
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer

# Download NLTK resources on first run
def download_nltk_resources():
    try:
        nltk.data.find('tokenizers/punkt')
        nltk.data.find('corpora/stopwords')
        nltk.data.find('corpora/wordnet')
    except LookupError:
        nltk.download('punkt')
        nltk.download('stopwords')
        nltk.download('wordnet')

# Call this at module import
download_nltk_resources()

def extract_keywords(text, additional_keywords=None):
    """Extract important keywords from job description."""
    # Tokenize and lowercase
    tokens = word_tokenize(text.lower())
    
    # Remove stopwords and punctuation
    stop_words = set(stopwords.words('english'))
    tokens = [token for token in tokens if token.isalnum() and token not in stop_words]
    
    # Lemmatize
    lemmatizer = WordNetLemmatizer()
    tokens = [lemmatizer.lemmatize(token) for token in tokens]
    
    # Count frequency
    freq_dist = nltk.FreqDist(tokens)
    
    # Get top keywords
    keywords = [word for word, freq in freq_dist.most_common(15)]
    
    # Add additional keywords if provided
    if additional_keywords:
        keywords.extend([k.lower() for k in additional_keywords])
    
    return list(set(keywords))  # Remove duplicates

def customize_cv(template_path, output_path, job_title, company, job_description, keywords=None):
    """Customize CV based on job description and keywords."""
    # Extract keywords from job description
    extracted_keywords = extract_keywords(job_description, keywords)
    
    # Load the template
    doc = Document(template_path)
    
    # Replace placeholders
    for paragraph in doc.paragraphs:
        paragraph.text = paragraph.text.replace('[JOB_TITLE]', job_title)
        paragraph.text = paragraph.text.replace('[COMPANY]', company)
    
    # Highlight skills based on keywords
    for paragraph in doc.paragraphs:
        for keyword in extracted_keywords:
            if keyword in paragraph.text.lower():
                # You could add formatting here, but for simplicity we'll just
                # ensure the keyword is in the document
                pass
    
    # Create temp docx file
    docx_path = output_path.replace('.pdf', '.docx')
    doc.save(docx_path)
    
    # Convert to PDF
    convert(docx_path, output_path)
    
    # Clean up the temporary docx file
    os.remove(docx_path)
    
    return output_path